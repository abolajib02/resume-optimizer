"""
Tests for the DOCX reconstructor.

Strategy: parse a template → reconstruct with various selection states →
re-parse the output and verify structure. This is a round-trip integration test.

Run: cd backend && pytest tests/test_reconstructor.py -v -s
"""
import os
from io import BytesIO

import pytest
from docx import Document

from app.models.request import DownloadRequest, FormattingSettings
from app.services.parser import parse_resume
from app.services.reconstructor import reconstruct_resume

CHRONOLOGICAL = os.path.join(
    os.path.expanduser("~"), "Downloads", "Chronological Resume Template.docx"
)
COMBINATION = os.path.join(
    os.path.expanduser("~"), "Downloads", "Combination Resume Template.docx"
)


def _load(path: str) -> bytes:
    with open(path, "rb") as f:
        return f.read()


def _open_docx(docx_bytes: bytes) -> Document:
    return Document(BytesIO(docx_bytes))


def _all_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ---------------------------------------------------------------------------
# Chronological round-trip
# ---------------------------------------------------------------------------

@pytest.mark.skipif(not os.path.exists(CHRONOLOGICAL), reason="Template not found")
class TestChronologicalRoundTrip:
    def setup_method(self):
        raw = _load(CHRONOLOGICAL)
        self.resume = parse_resume(raw)
        self.request = DownloadRequest(resume=self.resume)

    def test_produces_valid_docx_bytes(self):
        out = reconstruct_resume(self.request)
        assert isinstance(out, bytes)
        assert len(out) > 0
        # Must be a valid DOCX (ZIP) file
        assert out[:2] == b"PK"

    def test_output_is_openable(self):
        out = reconstruct_resume(self.request)
        doc = _open_docx(out)
        assert len(doc.paragraphs) > 0

    def test_candidate_name_present(self):
        out = reconstruct_resume(self.request)
        text = _all_text(_open_docx(out))
        assert self.resume.candidate_name in text

    def test_section_headings_present(self):
        out = reconstruct_resume(self.request)
        text = _all_text(_open_docx(out))
        for section in self.resume.sections:
            if section.heading:
                assert section.heading in text, (
                    f"Heading '{section.heading}' missing from output"
                )

    def test_job_titles_present(self):
        out = reconstruct_resume(self.request)
        text = _all_text(_open_docx(out))
        for section in self.resume.sections:
            for job in section.jobs:
                assert job.title in text, f"Job title '{job.title}' missing"

    def test_bullets_present(self):
        out = reconstruct_resume(self.request)
        text = _all_text(_open_docx(out))
        for section in self.resume.sections:
            for job in section.jobs:
                for bullet in job.bullets:
                    assert bullet.text in text, f"Bullet missing: {bullet.text[:40]}"

    def test_deselected_bullet_excluded(self):
        """Hiding a bullet should remove it from the output."""
        exp_section = next(
            s for s in self.resume.sections if s.section_type == "experience"
        )
        target_bullet = exp_section.jobs[0].bullets[0]
        request = DownloadRequest(
            resume=self.resume,
            selection={target_bullet.id: False},
        )
        out = reconstruct_resume(request)
        text = _all_text(_open_docx(out))
        assert target_bullet.text not in text

    def test_deselected_job_excluded(self):
        """Hiding a job should remove its header line and all bullets."""
        exp_section = next(
            s for s in self.resume.sections if s.section_type == "experience"
        )
        target_job = exp_section.jobs[0]
        # Deselect the job and all its bullets
        selection = {target_job.id: False}
        for b in target_job.bullets:
            selection[b.id] = False
        request = DownloadRequest(resume=self.resume, selection=selection)
        out = reconstruct_resume(request)
        doc = _open_docx(out)

        # The job header paragraph (title + date on same line) should be gone.
        # Use the full header line text rather than just the title to avoid
        # false matches in summary/skills paragraphs.
        header_line = f"{target_job.title}  {target_job.date_range}"
        para_texts = [p.text for p in doc.paragraphs]
        assert header_line not in para_texts, (
            f"Job header line '{header_line}' still present after deselection"
        )
        # All bullets should also be absent
        for bullet in target_job.bullets:
            assert bullet.text not in "\n".join(para_texts), (
                f"Bullet still present: {bullet.text[:50]}"
            )

    def test_inline_edit_applied(self):
        """An inline edit should replace the original bullet text."""
        exp_section = next(
            s for s in self.resume.sections if s.section_type == "experience"
        )
        target_bullet = exp_section.jobs[0].bullets[0]
        edited_text = "Custom edited bullet text for testing purposes"
        request = DownloadRequest(
            resume=self.resume,
            inline_edits={target_bullet.id: edited_text},
        )
        out = reconstruct_resume(request)
        text = _all_text(_open_docx(out))
        assert edited_text in text

    def test_page_dimensions_preserved(self):
        """Output document should have the same page size as the source."""
        out = reconstruct_resume(self.request)
        doc = _open_docx(out)
        emu = 914400
        width = round(doc.sections[0].page_width / emu, 1)
        height = round(doc.sections[0].page_height / emu, 1)
        assert width == 8.5
        assert height == 11.0

    def test_font_size_override(self):
        """All body runs should use the overridden font size."""
        request = DownloadRequest(
            resume=self.resume,
            settings=FormattingSettings(font_size_pt=10.0),
        )
        out = reconstruct_resume(request)
        doc = _open_docx(out)
        body_sizes = set()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size:
                    body_sizes.add(round(run.font.size / 12700, 1))
        # At least some runs should be at 10pt (headings may be larger)
        assert 10.0 in body_sizes, f"10pt not found in sizes: {body_sizes}"

    def test_section_order_override(self):
        """Re-ordering sections should change paragraph order in output."""
        sections = self.resume.sections
        if len(sections) < 2:
            pytest.skip("Not enough sections to reorder")

        # Reverse the section order
        reversed_order = [s.id for s in reversed(sections)]
        request = DownloadRequest(
            resume=self.resume,
            section_order=reversed_order,
        )
        out = reconstruct_resume(request)
        doc = _open_docx(out)
        text = _all_text(doc)

        # The last section's heading should appear before the first section's heading
        first_heading = sections[0].heading
        last_heading = sections[-1].heading
        if first_heading and last_heading:
            pos_first = text.find(first_heading)
            pos_last = text.find(last_heading)
            assert pos_last < pos_first, "Section reordering not reflected in output"

    def test_print_output_structure(self, capsys):
        """Print the reconstructed document paragraphs for inspection."""
        out = reconstruct_resume(self.request)
        doc = _open_docx(out)
        print(f"\n{'='*60}")
        print("Reconstructed paragraphs:")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                sizes = [
                    round(r.font.size / 12700, 1)
                    for r in para.runs
                    if r.font.size
                ]
                print(f"  [{i:3d}] style={para.style.name!r:20s} "
                      f"sizes={sizes} text={para.text[:60]!r}")
        print("="*60)


# ---------------------------------------------------------------------------
# Combination round-trip
# ---------------------------------------------------------------------------

@pytest.mark.skipif(not os.path.exists(COMBINATION), reason="Template not found")
class TestCombinationRoundTrip:
    def setup_method(self):
        raw = _load(COMBINATION)
        self.resume = parse_resume(raw)
        self.request = DownloadRequest(
            resume=self.resume,
            settings=FormattingSettings(format="combination"),
        )

    def test_produces_valid_docx(self):
        out = reconstruct_resume(self.request)
        assert out[:2] == b"PK"
        doc = _open_docx(out)
        assert len(doc.paragraphs) > 0

    def test_candidate_name_present(self):
        out = reconstruct_resume(self.request)
        text = _all_text(_open_docx(out))
        assert self.resume.candidate_name in text

    def test_all_sections_present(self):
        out = reconstruct_resume(self.request)
        text = _all_text(_open_docx(out))
        for section in self.resume.sections:
            if section.heading:
                assert section.heading in text

    def test_print_output_structure(self, capsys):
        out = reconstruct_resume(self.request)
        doc = _open_docx(out)
        print(f"\n{'='*60}")
        print("Combination reconstructed paragraphs:")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"  [{i:3d}] {para.text[:70]!r}")
        print("="*60)


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------

def test_empty_selection_produces_minimal_doc():
    """Deselecting everything should produce a document with no body content."""
    if not os.path.exists(CHRONOLOGICAL):
        pytest.skip("Template not found")
    resume = parse_resume(_load(CHRONOLOGICAL))
    # Deselect every node
    selection = {}
    for section in resume.sections:
        selection[section.id] = False
        for job in section.jobs:
            selection[job.id] = False
            for bullet in job.bullets:
                selection[bullet.id] = False
        for fp in section.free_paragraphs:
            selection[fp.id] = False
    request = DownloadRequest(resume=resume, selection=selection)
    out = reconstruct_resume(request)
    doc = _open_docx(out)
    # No visible text expected
    visible = [p.text for p in doc.paragraphs if p.text.strip()]
    assert visible == [], f"Expected empty doc, got: {visible}"


def test_margin_override():
    """Custom margins should be applied to the output document."""
    if not os.path.exists(CHRONOLOGICAL):
        pytest.skip("Template not found")
    resume = parse_resume(_load(CHRONOLOGICAL))
    request = DownloadRequest(
        resume=resume,
        settings=FormattingSettings(
            margins_inches={"top": 0.5, "bottom": 0.5, "left": 0.75, "right": 0.75}
        ),
    )
    out = reconstruct_resume(request)
    doc = _open_docx(out)
    emu = 914400
    top = round(doc.sections[0].top_margin / emu, 2)
    left = round(doc.sections[0].left_margin / emu, 2)
    assert top == 0.5, f"Top margin wrong: {top}"
    assert left == 0.75, f"Left margin wrong: {left}"
