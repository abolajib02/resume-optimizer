"""
DOCX Reconstructor
==================
Rebuilds a .docx file from a ParsedResume + SelectionState + FormattingSettings.

The reconstructor works entirely from serialized data (no original file needed),
making the API stateless. Run styles captured during parsing are replayed onto
a new document, and standard Word paragraph styles are applied where possible.

Key decisions:
  - Page dimensions and margins are taken from source_styles, then overridden
    by settings.margins_inches if the user changed them.
  - Font size override from settings applies to body text only. Heading-sized
    text is scaled proportionally to preserve the visual hierarchy.
  - For Combination format, skill_groups (from AnalysisResult) reorder bullets
    under category headings; Employment History is written in collapsed form.
  - Inline edits from the frontend replace bullet/paragraph text verbatim.
"""
from __future__ import annotations

from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph as DocxParagraph

from app.models.analysis import SkillGroup
from app.models.request import DownloadRequest, FormattingSettings
from app.models.resume import (
    Bullet,
    DocumentStyles,
    FreeParagraph,
    Job,
    ParsedResume,
    RunStyle,
    Section,
)

_EMU_PER_INCH = 914400


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def reconstruct_resume(request: DownloadRequest) -> bytes:
    """
    Build and return a DOCX file as bytes from the given DownloadRequest.
    """
    resume = request.resume
    settings = request.settings
    selection = request.selection
    inline_edits = request.inline_edits

    # Build lookup maps for ordering overrides sent by the frontend
    section_order = request.section_order or [s.id for s in resume.sections]
    job_order_map = request.job_order        # section_id → [job_id, ...]
    bullet_order_map = request.bullet_order  # job_id → [bullet_id, ...]

    doc = _setup_document(resume.source_styles, settings)

    section_map = {s.id: s for s in resume.sections}

    for section_id in section_order:
        section = section_map.get(section_id)
        if section is None:
            continue
        if not _is_visible(section_id, selection):
            continue

        _write_section(
            doc=doc,
            section=section,
            settings=settings,
            selection=selection,
            job_order=job_order_map.get(section_id, [j.id for j in section.jobs]),
            bullet_order_map=bullet_order_map,
            inline_edits=inline_edits,
            skill_groups=request.skill_groups,
        )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def _setup_document(source: DocumentStyles, settings: FormattingSettings) -> Document:
    doc = Document()

    # Remove the default empty paragraph that python-docx always adds
    for para in doc.paragraphs:
        _remove_paragraph(para)

    section = doc.sections[0]

    # Page size — use source dimensions (default Letter if unknown)
    w = source.page_width_inches or 8.5
    h = source.page_height_inches or 11.0
    section.page_width = int(w * _EMU_PER_INCH)
    section.page_height = int(h * _EMU_PER_INCH)

    # Margins — user override takes precedence over source
    m = settings.margins_inches
    section.top_margin = Inches(m.get("top", source.margin_top_inches or 1.0))
    section.bottom_margin = Inches(m.get("bottom", source.margin_bottom_inches or 1.0))
    section.left_margin = Inches(m.get("left", source.margin_left_inches or 1.0))
    section.right_margin = Inches(m.get("right", source.margin_right_inches or 1.0))

    return doc


# ---------------------------------------------------------------------------
# Section dispatcher
# ---------------------------------------------------------------------------

def _write_section(
    doc: Document,
    section: Section,
    settings: FormattingSettings,
    selection: dict[str, bool],
    job_order: list[str],
    bullet_order_map: dict[str, list[str]],
    inline_edits: dict[str, str],
    skill_groups: Optional[list[SkillGroup]],
) -> None:
    # Write the section heading (skip for implicit contact blocks with no heading)
    if section.heading:
        _write_paragraph_from_runs(
            doc=doc,
            text=section.heading,
            run_styles=section.heading_run_styles,
            paragraph_style=section.heading_paragraph_style,
            settings=settings,
            is_heading=True,
            inline_edit=inline_edits.get(section.id),
        )

    is_combination = settings.format == "combination"

    if section.section_type == "experience":
        _write_experience_section(
            doc=doc,
            section=section,
            settings=settings,
            selection=selection,
            job_order=job_order,
            bullet_order_map=bullet_order_map,
            inline_edits=inline_edits,
            collapsed=is_combination,
        )
    elif section.section_type == "skills" and is_combination and skill_groups:
        _write_combination_skills_section(
            doc=doc,
            section=section,
            settings=settings,
            selection=selection,
            skill_groups=skill_groups,
            inline_edits=inline_edits,
        )
    else:
        _write_free_section(
            doc=doc,
            section=section,
            settings=settings,
            selection=selection,
            inline_edits=inline_edits,
        )


# ---------------------------------------------------------------------------
# Experience section (chronological: full bullets; combination: collapsed)
# ---------------------------------------------------------------------------

def _write_experience_section(
    doc: Document,
    section: Section,
    settings: FormattingSettings,
    selection: dict[str, bool],
    job_order: list[str],
    bullet_order_map: dict[str, list[str]],
    inline_edits: dict[str, str],
    collapsed: bool,
) -> None:
    job_map = {j.id: j for j in section.jobs}

    for job_id in job_order:
        job = job_map.get(job_id)
        if job is None:
            continue
        if not _is_visible(job_id, selection):
            continue

        _write_job_header(doc, job, settings, inline_edits)

        if not collapsed:
            # Chronological: write selected bullets in user-defined order
            bullet_order = bullet_order_map.get(job_id, [b.id for b in job.bullets])
            bullet_map = {b.id: b for b in job.bullets}
            for bullet_id in bullet_order:
                bullet = bullet_map.get(bullet_id)
                if bullet is None:
                    continue
                if not _is_visible(bullet_id, selection):
                    continue
                _write_bullet(doc, bullet, settings, inline_edits)


def _write_job_header(
    doc: Document,
    job: Job,
    settings: FormattingSettings,
    inline_edits: dict[str, str],
) -> None:
    """
    Write the job title and company/date line(s).

    Template variants handled:
      A) Title and date are on the same paragraph (title_run_styles == date_run_styles)
         → write one paragraph: "{title}  {date_range}"
      B) Title is on its own paragraph, date is separate
         → write title paragraph, then date/company paragraph
    """
    title_text = inline_edits.get(job.id + "_title", job.title)
    company_text = inline_edits.get(job.id + "_company", job.company)
    date_text = inline_edits.get(job.id + "_date", job.date_range)

    # Detect variant A: title was extracted from the date line
    # (they share the same run_styles list identity — check by content equality)
    title_from_date_line = (
        job.title_run_styles
        and job.date_run_styles
        and job.title_run_styles == job.date_run_styles
    )

    if title_from_date_line:
        # Write as a single paragraph: "Title  date_range"
        # Build a merged run list: title text + spacer + date text
        combined_text = f"{title_text}  {date_text}"
        _write_paragraph_from_runs(
            doc=doc,
            text=combined_text,
            run_styles=job.title_run_styles,
            paragraph_style=job.title_paragraph_style,
            settings=settings,
            is_heading=False,
        )
    else:
        # Title on its own line
        if title_text:
            _write_paragraph_from_runs(
                doc=doc,
                text=title_text,
                run_styles=job.title_run_styles,
                paragraph_style=job.title_paragraph_style,
                settings=settings,
                is_heading=False,
            )
        # Company + location + date on the next line
        company_line_parts = [p for p in [company_text, job.location, date_text] if p]
        company_line = "\t".join(company_line_parts) if company_line_parts else ""
        if company_line:
            run_styles = job.company_run_styles or job.date_run_styles
            _write_paragraph_from_runs(
                doc=doc,
                text=company_line,
                run_styles=run_styles,
                paragraph_style=job.company_paragraph_style or job.date_paragraph_style,
                settings=settings,
                is_heading=False,
            )


def _write_bullet(
    doc: Document,
    bullet: Bullet,
    settings: FormattingSettings,
    inline_edits: dict[str, str],
) -> None:
    text = inline_edits.get(bullet.id, bullet.text)
    _write_paragraph_from_runs(
        doc=doc,
        text=text,
        run_styles=bullet.run_styles,
        paragraph_style=bullet.paragraph_style,
        settings=settings,
        is_heading=False,
    )


# ---------------------------------------------------------------------------
# Combination-format skills section (LLM-grouped categories)
# ---------------------------------------------------------------------------

def _write_combination_skills_section(
    doc: Document,
    section: Section,
    settings: FormattingSettings,
    selection: dict[str, bool],
    skill_groups: list[SkillGroup],
    inline_edits: dict[str, str],
) -> None:
    """
    Write skill bullets organized under LLM-generated category headings.
    Each SkillGroup.heading is written as a sub-heading (bold, same font as body).
    Only bullets that are both in the skill group and in the selection are written.
    """
    # Build a lookup: bullet_id → FreeParagraph
    fp_map = {fp.id: fp for fp in section.free_paragraphs}

    for group in skill_groups:
        # Write category heading — styled as bold body text (not a Word Heading)
        _write_plain_paragraph(
            doc=doc,
            text=group.heading,
            font_name=_fallback_font(section),
            font_size_pt=settings.font_size_pt,
            bold=True,
            paragraph_style="Normal",
        )
        for bullet_id in group.bullet_ids:
            if not _is_visible(bullet_id, selection):
                continue
            fp = fp_map.get(bullet_id)
            if fp is None:
                continue
            text = inline_edits.get(fp.id, fp.text)
            _write_paragraph_from_runs(
                doc=doc,
                text=text,
                run_styles=fp.run_styles,
                paragraph_style=fp.paragraph_style,
                settings=settings,
                is_heading=False,
            )


# ---------------------------------------------------------------------------
# Free-paragraph sections (skills, education, summary, etc.)
# ---------------------------------------------------------------------------

def _write_free_section(
    doc: Document,
    section: Section,
    settings: FormattingSettings,
    selection: dict[str, bool],
    inline_edits: dict[str, str],
) -> None:
    for fp in section.free_paragraphs:
        if not _is_visible(fp.id, selection):
            continue
        text = inline_edits.get(fp.id, fp.text)
        _write_paragraph_from_runs(
            doc=doc,
            text=text,
            run_styles=fp.run_styles,
            paragraph_style=fp.paragraph_style,
            settings=settings,
            is_heading=False,
        )


# ---------------------------------------------------------------------------
# Core paragraph writers
# ---------------------------------------------------------------------------

def _write_paragraph_from_runs(
    doc: Document,
    text: str,
    run_styles: list[RunStyle],
    paragraph_style: str,
    settings: FormattingSettings,
    is_heading: bool,
    inline_edit: Optional[str] = None,
) -> DocxParagraph:
    """
    Add a paragraph to `doc`, replaying stored run styles.

    If run_styles is empty or doesn't cover the full text, the text is written
    as a single run with style inferred from settings.

    Font size logic:
      - Body text: use settings.font_size_pt (user override)
      - Heading text: use the stored size (preserves visual hierarchy), but
        clamp to a minimum of settings.font_size_pt so headings never shrink
        below body size.
    """
    display_text = inline_edit if inline_edit is not None else text

    # Try to apply the stored paragraph style; fall back to Normal
    style_name = _safe_style(doc, paragraph_style)
    para = doc.add_paragraph(style=style_name)

    if not run_styles:
        # No run data — write as a single plain run
        run = para.add_run(display_text)
        _apply_size(run, settings.font_size_pt, is_heading, None)
        return para

    # Check if the stored runs collectively reconstruct the original text.
    # If they do, replay them. Otherwise fall back to single-run.
    stored_text = "".join(r.text for r in run_styles)

    if stored_text == text or not text:
        # Faithful replay
        for rs in run_styles:
            run_text = rs.text
            # If this is an inline-edited paragraph, distribute the edit across
            # the first run only and zero out the rest (simple but effective).
            run = para.add_run(run_text)
            _apply_run_style(run, rs, settings.font_size_pt, is_heading)
        # If the user edited the text, replace everything with a single run
        if inline_edit is not None and inline_edit != text:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = inline_edit
            else:
                run = para.add_run(inline_edit)
                _apply_size(run, settings.font_size_pt, is_heading, None)
    else:
        # Stored runs don't match — write the text as one run with inferred style
        run = para.add_run(display_text)
        # Inherit font from the first run style if available
        if run_styles:
            rs = run_styles[0]
            _apply_run_style(run, rs, settings.font_size_pt, is_heading)
        else:
            _apply_size(run, settings.font_size_pt, is_heading, None)

    return para


def _write_plain_paragraph(
    doc: Document,
    text: str,
    font_name: Optional[str],
    font_size_pt: float,
    bold: bool = False,
    paragraph_style: str = "Normal",
) -> DocxParagraph:
    """Write a paragraph with uniform formatting (no run_styles replay)."""
    style_name = _safe_style(doc, paragraph_style)
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(text)
    run.bold = bold
    if font_name:
        run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    return para


# ---------------------------------------------------------------------------
# Run-level style application
# ---------------------------------------------------------------------------

def _apply_run_style(
    run,
    rs: RunStyle,
    body_font_size_pt: float,
    is_heading: bool,
) -> None:
    if rs.bold is not None:
        run.bold = rs.bold
    if rs.italic is not None:
        run.italic = rs.italic
    if rs.underline is not None:
        run.underline = rs.underline
    if rs.font_name:
        run.font.name = rs.font_name
    _apply_size(run, body_font_size_pt, is_heading, rs.font_size_pt)
    if rs.color_hex:
        try:
            run.font.color.rgb = RGBColor.from_string(rs.color_hex)
        except Exception:
            pass


def _apply_size(
    run,
    body_size: float,
    is_heading: bool,
    stored_size: Optional[float],
) -> None:
    """
    Determine the effective font size for a run:
      - Heading: use stored_size if available (preserves hierarchy), clamped
        to at least body_size.
      - Body: always use body_size (the user's chosen value).
    """
    if is_heading and stored_size:
        effective = max(stored_size, body_size)
    else:
        effective = body_size
    run.font.size = Pt(effective)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _is_visible(node_id: str, selection: dict[str, bool]) -> bool:
    """Return True if the node is selected (missing from map = visible by default)."""
    return selection.get(node_id, True)


def _safe_style(doc: Document, style_name: str) -> str:
    """Return style_name if it exists in the document, else 'Normal'."""
    try:
        doc.styles[style_name]
        return style_name
    except KeyError:
        return "Normal"


def _fallback_font(section: Section) -> Optional[str]:
    """Return the first non-None font name found in the section's free_paragraphs."""
    for fp in section.free_paragraphs:
        for rs in fp.run_styles:
            if rs.font_name:
                return rs.font_name
    return None


def _remove_paragraph(para: DocxParagraph) -> None:
    """Remove a paragraph element from the document body."""
    p = para._element
    p.getparent().remove(p)
