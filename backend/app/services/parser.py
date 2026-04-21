"""
DOCX Resume Parser
==================
Converts an uploaded .docx file into a structured ParsedResume tree.

Pipeline:
  Stage 1  – Load document, capture style defaults and page layout
  Stage 2  – Detect if the document uses multi-column tables (flag as warning)
  Stage 3  – Iterate all body paragraphs and classify each one:
               heading | bullet | date-line | free-paragraph
  Stage 4  – Group classified paragraphs into Sections, Jobs, and Bullets
  Stage 5  – Detect resume format (chronological vs. combination)
  Stage 6  – Extract candidate name from the first heading/paragraph
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.text.paragraph import Paragraph

from app.models.resume import (
    Bullet,
    DocumentStyles,
    FreeParagraph,
    Job,
    ParsedResume,
    RunStyle,
    Section,
)
from app.utils.docx_helpers import (
    BULLET_CHARS,
    classify_section,
    contains_date_range,
    extract_date_range,
    paragraph_font_size_pt,
    paragraph_has_bottom_border,
    paragraph_is_bold,
    paragraph_is_underlined,
    paragraph_left_indent_emu,
    serialize_runs,
    strip_bullet_char,
)

# ---------------------------------------------------------------------------
# Internal classification labels used during the parse pass
# ---------------------------------------------------------------------------
_CLS_HEADING = "heading"
_CLS_BULLET = "bullet"
_CLS_DATE_LINE = "date_line"  # line that contains a date range
_CLS_CONTACT = "contact"
_CLS_FREE = "free"  # everything else
_CLS_EMPTY = "empty"


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def parse_resume(file_bytes: bytes) -> ParsedResume:
    """Parse a DOCX file (as raw bytes) into a ParsedResume."""
    doc = Document(BytesIO(file_bytes))
    warnings: list[str] = []

    # Stage 1 – document-level style defaults
    source_styles = _extract_document_styles(doc)

    # Stage 2 – table detection
    if _has_multi_column_tables(doc):
        warnings.append(
            "This resume uses tables or text boxes for layout. "
            "Some content may not parse correctly. "
            "Consider uploading a single-column version for best results."
        )

    # Collect all body paragraphs (tables linearized into paragraphs)
    all_paragraphs = _collect_paragraphs(doc, warnings)

    # Compute the document's body font size for heading detection
    body_font_size = _estimate_body_font_size(all_paragraphs)

    # Stage 3 – classify each paragraph
    classified = [(para, _classify_paragraph(para, body_font_size)) for para in all_paragraphs]

    # Stage 4 – build section/job/bullet tree
    sections = _build_sections(classified, warnings)

    # Stage 5 – detect format
    detected_format = _detect_format(sections)

    # Stage 6 – extract candidate name
    candidate_name = _extract_candidate_name(classified)

    return ParsedResume(
        candidate_name=candidate_name,
        sections=sections,
        detected_format=detected_format,
        source_styles=source_styles,
        parse_warnings=warnings,
    )


# ---------------------------------------------------------------------------
# Stage 1 – Document-level style extraction
# ---------------------------------------------------------------------------


def _extract_document_styles(doc: Document) -> DocumentStyles:
    styles = DocumentStyles()

    # Page layout from the first section
    try:
        sec = doc.sections[0]
        emu = 914400  # 1 inch in EMU
        styles.page_width_inches = round(sec.page_width / emu, 3) if sec.page_width else None
        styles.page_height_inches = round(sec.page_height / emu, 3) if sec.page_height else None
        styles.margin_top_inches = round(sec.top_margin / emu, 3) if sec.top_margin else None
        styles.margin_bottom_inches = (
            round(sec.bottom_margin / emu, 3) if sec.bottom_margin else None
        )
        styles.margin_left_inches = round(sec.left_margin / emu, 3) if sec.left_margin else None
        styles.margin_right_inches = round(sec.right_margin / emu, 3) if sec.right_margin else None
    except Exception:
        pass

    # Default font from document styles
    try:
        default_style = doc.styles["Normal"]
        if default_style.font.name:
            styles.default_font_name = default_style.font.name
        if default_style.font.size:
            styles.default_font_size_pt = round(default_style.font.size / 12700, 2)
    except Exception:
        pass

    return styles


# ---------------------------------------------------------------------------
# Stage 2 – Table detection
# ---------------------------------------------------------------------------


def _has_multi_column_tables(doc: Document) -> bool:
    """Return True if the document body contains tables (possible multi-column layout)."""
    return len(doc.tables) > 0


# ---------------------------------------------------------------------------
# Paragraph collection (including table cells)
# ---------------------------------------------------------------------------


def _collect_paragraphs(doc: Document, warnings: list[str]) -> list[Paragraph]:
    """
    Return all paragraphs in document order.
    Table cells are linearized: row by row, cell by cell.
    """
    paragraphs: list[Paragraph] = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
        if tag == "p":
            # Wrap back into a Paragraph object
            from docx.text.paragraph import Paragraph as P

            paragraphs.append(P(block, doc))
        elif tag == "tbl":
            # Linearize table cells
            from docx.table import Table

            tbl = Table(block, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        paragraphs.append(para)

    return paragraphs


# ---------------------------------------------------------------------------
# Body font size estimation
# ---------------------------------------------------------------------------


def _estimate_body_font_size(paragraphs: list[Paragraph]) -> float:
    """
    Estimate the predominant body font size by taking the mode of all
    non-None font sizes across paragraphs, weighted toward 11pt.
    """
    sizes: list[float] = []
    for para in paragraphs:
        sz = paragraph_font_size_pt(para)
        if sz and 6 <= sz <= 14:
            sizes.append(sz)
    if not sizes:
        return 11.0
    return max(set(sizes), key=sizes.count)


# ---------------------------------------------------------------------------
# Stage 3 – Paragraph classification
# ---------------------------------------------------------------------------


def _classify_paragraph(para: Paragraph, body_font_size: float) -> str:
    text = para.text.strip()

    if not text:
        return _CLS_EMPTY

    style_name = para.style.name if para.style else ""
    font_size = paragraph_font_size_pt(para)

    # ---- Heading signals ----
    heading_signals = 0

    # 1. Word style is a heading
    if style_name.lower().startswith("heading"):
        heading_signals += 3  # strong signal

    # 2. All-caps, short, no punctuation (except ampersand/slash)
    clean = re.sub(r"[&/\-]", "", text)
    if (
        clean == clean.upper()
        and len(text.split()) <= 6
        and not re.search(r"[.@,!?]", text)
        and len(text) > 2
    ):
        heading_signals += 2

    # 3. Font size notably larger than body
    if font_size and font_size >= body_font_size + 1.5:
        heading_signals += 2

    # 4. Bold and short — +1 only; not sufficient alone so that bold-only
    #    content lines (project titles, job sub-headings like "Technical Skills:")
    #    don't get mis-classified as section headings.
    if paragraph_is_bold(para) and len(text) < 40 and len(text.split()) <= 5:
        heading_signals += 1

    # 5. Has a bottom border (decorative section divider)
    if paragraph_has_bottom_border(para):
        heading_signals += 2

    # 6. Underlined text — resumes commonly underline true section headings
    #    (e.g. "Projects", "Professional Experience") but not inline sub-headings
    #    or content bold lines. Combined with signal 4 this gives bold+underlined
    #    paragraphs a total of 3, crossing the threshold cleanly.
    if paragraph_is_underlined(para):
        heading_signals += 2

    if heading_signals >= 2:
        return _CLS_HEADING

    # ---- Bullet signals ----
    # 1. Style name suggests a list
    if any(kw in style_name.lower() for kw in ("list", "bullet", "item")):
        return _CLS_BULLET

    # 2. Starts with a known bullet character
    if text and text[0] in BULLET_CHARS:
        return _CLS_BULLET

    # 3. Has a meaningful left indent relative to body
    indent = paragraph_left_indent_emu(para)
    if indent and indent > 200000:  # ~0.22 inches
        return _CLS_BULLET

    # ---- Date line ----
    if contains_date_range(text):
        return _CLS_DATE_LINE

    return _CLS_FREE


# ---------------------------------------------------------------------------
# Stage 4 – Build section/job/bullet tree
# ---------------------------------------------------------------------------


def _build_sections(
    classified: list[tuple[Paragraph, str]],
    warnings: list[str],
) -> list[Section]:
    """
    Walk the classified paragraph list and assemble it into Sections.
    Within experience sections, further group paragraphs into Jobs.
    """
    sections: list[Section] = []
    current_section: Section | None = None
    pending_free: list[tuple[Paragraph, int]] = []  # paragraphs before first heading
    section_order = 0

    i = 0
    while i < len(classified):
        para, cls = classified[i]
        text = para.text.strip()

        if cls == _CLS_EMPTY:
            i += 1
            continue

        if cls == _CLS_HEADING:
            # Save any accumulated free paragraphs into the previous section
            if current_section is not None and pending_free:
                _flush_free_paragraphs(current_section, pending_free)
                pending_free = []

            section_type = classify_section(text)

            current_section = Section(
                heading=text,
                heading_run_styles=[RunStyle(**r) for r in serialize_runs(para)],
                heading_paragraph_style=para.style.name if para.style else "Normal",
                section_type=section_type,
                original_order=section_order,
            )
            sections.append(current_section)
            section_order += 1
            i += 1
            continue

        # No section yet — create an implicit contact/header section
        if current_section is None:
            current_section = Section(
                heading="",
                section_type="contact",
                original_order=section_order,
            )
            sections.append(current_section)
            section_order += 1

        if current_section.section_type == "experience":
            # Delegate experience parsing to a sub-routine
            consumed = _parse_experience_block(classified, i, current_section, warnings)
            i += consumed
        else:
            # Non-experience section: accumulate bullets and free paragraphs
            fp_index = len(current_section.free_paragraphs)
            if cls == _CLS_BULLET:
                bullet = Bullet(
                    text=strip_bullet_char(text),
                    run_styles=[RunStyle(**r) for r in serialize_runs(para)],
                    paragraph_style=para.style.name if para.style else "Normal",
                    original_index=fp_index,
                )
                current_section.free_paragraphs.append(
                    FreeParagraph(
                        text=bullet.text,
                        run_styles=bullet.run_styles,
                        paragraph_style=bullet.paragraph_style,
                        original_index=fp_index,
                    )
                )
            else:
                fp = FreeParagraph(
                    text=text,
                    run_styles=[RunStyle(**r) for r in serialize_runs(para)],
                    paragraph_style=para.style.name if para.style else "Normal",
                    original_index=fp_index,
                )
                current_section.free_paragraphs.append(fp)
            i += 1

    return sections


def _flush_free_paragraphs(section: Section, pending: list[tuple[Paragraph, int]]) -> None:
    for para, idx in pending:
        text = para.text.strip()
        if not text:
            continue
        section.free_paragraphs.append(
            FreeParagraph(
                text=text,
                run_styles=[RunStyle(**r) for r in serialize_runs(para)],
                paragraph_style=para.style.name if para.style else "Normal",
                original_index=idx,
            )
        )


def _parse_experience_block(
    classified: list[tuple[Paragraph, str]],
    start: int,
    section: Section,
    warnings: list[str],
) -> int:
    """
    Parse one or more Job blocks starting at `start` within an experience section.
    Returns the number of paragraphs consumed.
    """
    i = start
    job_order = len(section.jobs)

    # Buffer for lines between the section heading and the first date line.
    # These form the job header (title, company, location).
    header_buffer: list[tuple[Paragraph, str]] = []
    current_job: Job | None = None
    bullet_index = 0

    while i < len(classified):
        para, cls = classified[i]
        text = para.text.strip()

        # Stop when we hit the next section heading
        if cls == _CLS_HEADING:
            break

        if cls == _CLS_EMPTY:
            i += 1
            continue

        if cls == _CLS_DATE_LINE:
            # A date line signals the start of a new job block.
            # Finalize any previous job first.
            if current_job is not None:
                section.jobs.append(current_job)

            # Build the new job from header_buffer + this date line
            current_job = _build_job(header_buffer, para, text, job_order, warnings)
            job_order += 1
            bullet_index = 0
            header_buffer = []
            i += 1

            # Peek ahead: if the job still has no company, the very next
            # non-empty, non-bullet, non-date paragraph is likely the company line.
            if not current_job.company and i < len(classified):
                peek_para, peek_cls = classified[i]
                peek_text = peek_para.text.strip()
                if peek_text and peek_cls not in (
                    _CLS_HEADING,
                    _CLS_BULLET,
                    _CLS_DATE_LINE,
                    _CLS_EMPTY,
                ):
                    # Split on tab in case company and location are tab-separated
                    # e.g. "Minnesota Diversified Industries\tSt. Paul, MN"
                    parts = peek_text.split("\t", 1)
                    current_job.company = parts[0].strip()
                    current_job.company_run_styles = [
                        RunStyle(**r) for r in serialize_runs(peek_para)
                    ]
                    current_job.company_paragraph_style = (
                        peek_para.style.name if peek_para.style else "Normal"
                    )
                    if len(parts) > 1 and not current_job.location:
                        current_job.location = parts[1].strip()
                    i += 1
            continue

        if cls == _CLS_BULLET and current_job is not None:
            bullet = Bullet(
                text=strip_bullet_char(text),
                run_styles=[RunStyle(**r) for r in serialize_runs(para)],
                paragraph_style=para.style.name if para.style else "Normal",
                original_index=bullet_index,
            )
            current_job.bullets.append(bullet)
            bullet_index += 1
            i += 1
            continue

        if current_job is None:
            # We're still in the job header area
            header_buffer.append((para, cls))
        else:
            # Free paragraph after job header but before bullets, or between bullets
            # (e.g., a sub-heading like "Key Achievements")
            fp = FreeParagraph(
                text=text,
                run_styles=[RunStyle(**r) for r in serialize_runs(para)],
                paragraph_style=para.style.name if para.style else "Normal",
                original_index=len(current_job.raw_header_paragraphs),
            )
            current_job.raw_header_paragraphs.append(fp)

        i += 1

    # Flush the last job
    if current_job is not None:
        section.jobs.append(current_job)
    elif header_buffer:
        # Experience section with no date lines — store as free paragraphs
        warnings.append(
            f"Could not detect job entries in section '{section.heading}'. "
            "Content stored as free paragraphs."
        )
        for para, _ in header_buffer:
            text = para.text.strip()
            if text:
                section.free_paragraphs.append(
                    FreeParagraph(
                        text=text,
                        run_styles=[RunStyle(**r) for r in serialize_runs(para)],
                        paragraph_style=para.style.name if para.style else "Normal",
                        original_index=len(section.free_paragraphs),
                    )
                )

    return i - start


def _build_job(
    header_buffer: list[tuple[Paragraph, str]],
    date_para: Paragraph,
    date_text: str,
    job_order: int,
    warnings: list[str],
) -> Job:
    """
    Construct a Job from buffered header paragraphs and the date-line paragraph.

    Assignment heuristics:
      1. header_buffer has 2+ items → first = title, second = company, rest = raw
      2. header_buffer has 1 item   → title only; company will be filled from the
                                       paragraph immediately following the date line
                                       (handled by the caller peeking ahead)
      3. header_buffer is empty     → title is extracted from the date line text
                                       by stripping the date range portion;
                                       company filled by caller peek

    Date range and any remaining text on the date line become date_range/location.
    """
    date_range = extract_date_range(date_text) or date_text
    date_run_styles = [RunStyle(**r) for r in serialize_runs(date_para)]
    date_paragraph_style = date_para.style.name if date_para.style else "Normal"

    title = ""
    title_run_styles: list[RunStyle] = []
    title_paragraph_style = "Normal"
    company = ""
    raw_extras: list[FreeParagraph] = []

    for idx, (para, _) in enumerate(header_buffer):
        t = para.text.strip()
        if idx == 0:
            title = t
            title_run_styles = [RunStyle(**r) for r in serialize_runs(para)]
            title_paragraph_style = para.style.name if para.style else "Normal"
        elif idx == 1:
            company = t
        else:
            raw_extras.append(
                FreeParagraph(
                    text=t,
                    run_styles=[RunStyle(**r) for r in serialize_runs(para)],
                    paragraph_style=para.style.name if para.style else "Normal",
                    original_index=idx,
                )
            )

    if not title:
        # Case: title is embedded on the same line as the date range.
        # Strip the date range and any separators to recover the title text.
        candidate = date_text.replace(date_range, "").strip(" ,|–—-\t")
        if candidate:
            title = candidate
            # Title lives on the date line — borrow its run styles
            title_run_styles = date_run_styles
            title_paragraph_style = date_paragraph_style
        else:
            # Single-line full entry: "Software Engineer, Acme Corp, 2020-2023"
            parts = re.split(r",\s*|\s+at\s+|\s+@\s+", date_text, maxsplit=2)
            if len(parts) >= 2:
                title = parts[0].strip()
                company = parts[1].strip()
                title_run_styles = date_run_styles
                title_paragraph_style = date_paragraph_style
                warnings.append(
                    f"Job header '{date_text[:60]}' parsed from single line — "
                    "please verify title and company."
                )

    # Location: anything left on the date line after removing date range and title
    remainder = date_text.replace(date_range, "")
    if title and title not in date_range:
        remainder = remainder.replace(title, "")
    location_candidate = remainder.strip(" ,|–—-\t")
    location = location_candidate if location_candidate else None

    return Job(
        title=title,
        title_run_styles=title_run_styles,
        title_paragraph_style=title_paragraph_style,
        date_range=date_range,
        date_run_styles=date_run_styles,
        date_paragraph_style=date_paragraph_style,
        company=company,
        location=location,
        raw_header_paragraphs=raw_extras,
        original_order=job_order,
    )


# ---------------------------------------------------------------------------
# Stage 5 – Format detection
# ---------------------------------------------------------------------------


def _detect_format(sections: list[Section]) -> str:
    """
    Heuristic: if a skills section appears before the experience section,
    and experience jobs have no/few bullets, it's likely Combination format.
    Otherwise Chronological.
    """
    section_types = [s.section_type for s in sections if s.section_type != "contact"]

    has_skills = "skills" in section_types
    has_experience = "experience" in section_types

    if not has_experience:
        return "unknown"

    skills_idx = section_types.index("skills") if has_skills else None
    exp_idx = section_types.index("experience")

    if skills_idx is not None and skills_idx < exp_idx:
        # Skills leads — check if experience bullets are sparse
        exp_section = next(s for s in sections if s.section_type == "experience")
        total_bullets = sum(len(j.bullets) for j in exp_section.jobs)
        if total_bullets <= len(exp_section.jobs) * 2:
            return "combination"

    return "chronological"


# ---------------------------------------------------------------------------
# Stage 6 – Candidate name extraction
# ---------------------------------------------------------------------------


def _extract_candidate_name(classified: list[tuple[Paragraph, str]]) -> str:
    """
    The candidate name is almost always the very first non-empty paragraph
    in the document, typically styled as a large heading.
    """
    for para, cls in classified:
        text = para.text.strip()
        if text and cls in (_CLS_HEADING, _CLS_FREE):
            # Sanity check: a name shouldn't contain @ (email) or digits (phone)
            if "@" not in text and not re.search(r"\d{5,}", text):
                return text
    return ""
